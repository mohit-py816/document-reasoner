import fitz  # PyMuPDF
from docx import Document
import xml.etree.ElementTree as ET
import yaml
import logging
from pathlib import Path

logger = logging.getLogger(__name__)

class DocumentParser:
    @staticmethod
    def parse(file_path: str) -> str:
        """Main parsing entry point"""
        path = Path(file_path)
        try:
            if path.suffix.lower() == '.pdf':
                return DocumentParser.parse_pdf(file_path)
            elif path.suffix.lower() == '.docx':
                return DocumentParser.parse_docx(file_path)
            elif path.suffix.lower() == '.xml':
                return DocumentParser.parse_xml(file_path)
            elif path.suffix.lower() in ('.yml', '.yaml'):
                return DocumentParser.parse_yaml(file_path)
            else:
                logger.error(f"Unsupported file type: {path.suffix}")
                return ""
        except Exception as e:
            logger.error(f"Error parsing {file_path}: {str(e)}")
            return ""
        if not Path(file_path).exists():
            logger.error(f"File not found: {file_path}")
            return ""


    @staticmethod
    def parse_pdf(file_path: str) -> str:
        """Extract text from PDF with formatting preservation"""
        text = []
        try:
            with fitz.open(file_path) as doc:
                for page in doc:
                    text.append(page.get_text("text"))
            return "\n".join(text)
        except Exception as e:
            logger.error(f"PDF parsing failed: {str(e)}")
            return ""
        for page in doc:
            text.append(f"--- Page {page.number + 1} ---\n{page.get_text('text')}")


    @staticmethod
    def parse_docx(file_path: str) -> str:
        """Extract text from DOCX documents"""
        try:
            doc = Document(file_path)
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            logger.error(f"DOCX parsing failed: {str(e)}")
            return ""
        text = []
        for i, para in enumerate(doc.paragraphs):
            if para.text.strip():
                text.append(f"Paragraph {i+1}: {para.text}")
        return "\n".join(text)


    @staticmethod
    def parse_xml(file_path: str) -> str:
        """Convert XML to human-readable text"""
        try:
            tree = ET.parse(file_path)
            root = tree.getroot()

            def process_element(element, depth=0):
                text = []
                if element.text and element.text.strip():
                    text.append("  " * depth + element.text.strip())
                for child in element:
                    text.append(process_element(child, depth + 1))
                return "\n".join(text)

            return process_element(root)
        except Exception as e:
            logger.error(f"XML parsing failed: {str(e)}")
            return ""

    @staticmethod
    def parse_yaml(file_path: str) -> str:
        """Convert YAML to formatted text"""
        try:
            with open(file_path, 'r') as f:
                data = yaml.safe_load(f)
            return yaml.dump(data, default_flow_style=False)
        except Exception as e:
            logger.error(f"YAML parsing failed: {str(e)}")
            return ""
